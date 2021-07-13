import docx
from docx.shared import Pt, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.text import WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
from docx.text.tabstops import WD_TAB_ALIGNMENT
from docx.text.tabstops import WD_TAB_LEADER
import json

with open('config.json', encoding='utf-8') as json_data_file:
    config = json.load(json_data_file)

document = docx.Document()

# Define default margin
sections = document.sections
for section in sections:
    section.top_margin = Cm(3)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(3)
    section.right_margin = Cm(2)


# Normal font to the file
font = document.styles['Normal'].font
font.name = 'Arial'
font.size = Pt(12)

pageCount = 3


def Cover():

    # Institution name capital and centralized
    inst = document.add_paragraph()
    inst.alignment = WD_ALIGN_PARAGRAPH.CENTER
    runInst = inst.add_run(config['institution'].upper())

    # Your name capital, bold and centralized
    name = document.add_paragraph()
    name.alignment = WD_ALIGN_PARAGRAPH.CENTER
    runName = name.add_run(config['yourname'].upper())
    runName.bold = True

    # Space to the title
    for line in range(9):
        document.add_paragraph()

    # Title capital, bold, centralized e bigger
    title = document.add_paragraph(config['title'].upper())
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.style = document.styles.add_style(
        'TituloABNT', WD_STYLE_TYPE.PARAGRAPH)
    fontTitle = title.style.font
    fontTitle.name = 'Arial'
    fontTitle.size = Pt(14)
    fontTitle.bold = True

    # Subtitle bold and centralized
    sub = document.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    runSub = sub.add_run(config['subtitle'])
    runSub.bold = True

    # Space to the city
    for lines in range(11 - config['titlelines']):
        document.add_paragraph()

    # City capital and centralized
    city = document.add_paragraph()
    city.alignment = WD_ALIGN_PARAGRAPH.CENTER
    runCity = city.add_run(config['city'].upper())

    # Year capital and centralized
    year = document.add_paragraph()
    year.alignment = WD_ALIGN_PARAGRAPH.CENTER
    runYear = year.add_run(config['year'])


def BackCover():

    # Your name capital, bold and centralized
    name = document.add_paragraph()
    name.alignment = WD_ALIGN_PARAGRAPH.CENTER
    runName = name.add_run(config['yourname'].upper())
    runName.bold = True

    # Space to the title
    for line in range(7):
        document.add_paragraph()

    # Title capital, bold, centralized e bigger
    title = document.add_paragraph(
        text=config['title'].upper(), style='TituloABNT')
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Subtitle bold and centralized
    sub = document.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    runSub = sub.add_run(config['subtitle'])
    runSub.bold = True

    # Space to the indicative note
    for line in range(4):
        document.add_paragraph()

    # Indicative note right aligned, smaller and indented
    note = document.add_paragraph(config['notetext'])
    note.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    note.style = document.styles.add_style(
        'IndicativeNote', WD_STYLE_TYPE.PARAGRAPH)
    note.paragraph_format.left_indent = Cm(9)
    fontNote = note.style.font
    fontNote.name = 'Arial'
    fontNote.size = Pt(10)

    # Advisor in the note style
    advisor = document.add_paragraph(
        text='Orientador(a): '+config['advisor'], style='IndicativeNote')
    advisor.paragraph_format.left_indent = Cm(9)

    for lines in range(4):
        lineSpacing = document.add_paragraph(style='IndicativeNote')
        lineSpacing.paragraph_format.space_before = Pt(0)
        lineSpacing.paragraph_format.space_after = Pt(0)

    # Space to the city
    for lines in range(5 - config['titlelines']):
        document.add_paragraph()

    # City capital and centralized
    city = document.add_paragraph()
    city.alignment = WD_ALIGN_PARAGRAPH.CENTER
    runCity = city.add_run(config['city'].upper())

    # Year bold and centralized
    year = document.add_paragraph()
    year.alignment = WD_ALIGN_PARAGRAPH.CENTER
    runYear = year.add_run(config['year'])


def Dedicatory():
    global pageCount
    if config['dedicatory'] == 1:
        document.add_paragraph()

        dedication = document.add_paragraph("Write here...")
        dedication.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        dedication.paragraph_format.space_before = Pt(475)
        dedication.style = document.styles.add_style(
            'Dedicatory', WD_STYLE_TYPE.PARAGRAPH)
        dedication.paragraph_format.left_indent = Cm(9)
        fontDedication = dedication.style.font
        fontDedication.name = 'Arial'
        fontDedication.size = Pt(12)

        pageCount += 1
        document.add_page_break()


def Thanks():
    global pageCount
    if config['thanks'] == 1:
        title = document.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        title.paragraph_format.space_before = 0
        title.paragraph_format.space_after = 0
        runThanks = title.add_run("AGRADECIMENTOS\n")
        runThanks.bold = True

        thanksText = document.add_paragraph(
            "Write here...")
        thanksText.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        thanksText.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        thanksText.paragraph_format.space_before = 0
        thanksText.paragraph_format.space_after = 0
        thanksText.paragraph_format.first_line_indent = Cm(1.25)

        pageCount += 1
        document.add_page_break()


def Epigraph():
    global pageCount
    if config['epigraph'] == 1:
        document.add_paragraph()

        epigraph = document.add_paragraph(
            "Write here...")
        epigraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        epigraph.paragraph_format.space_before = Pt(475)
        epigraph.style = document.styles.add_style(
            'Epigraph', WD_STYLE_TYPE.PARAGRAPH)
        epigraph.paragraph_format.left_indent = Cm(9)
        fontepigraph = epigraph.style.font
        fontepigraph.name = 'Arial'
        fontepigraph.size = Pt(12)

        author = document.add_paragraph("Name of author")
        author.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        pageCount += 1
        document.add_page_break()


def Resume():
    global pageCount
    if config['resume'] == 1:
        title = document.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        title.paragraph_format.space_before = 0
        title.paragraph_format.space_after = 0
        runResume = title.add_run("RESUMO\n")
        runResume.bold = True

        resumeText = document.add_paragraph(
            "Write here...")
        resumeText.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        resumeText.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        resumeText.paragraph_format.space_before = 0
        resumeText.paragraph_format.space_after = 0
        resumeText.paragraph_format.first_line_indent = Cm(1.25)

        pageCount += 1
        document.add_page_break()

        if config['abstract'] == 1:
            title = document.add_paragraph()
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            title.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
            title.paragraph_format.space_before = 0
            title.paragraph_format.space_after = 0
            runAbstract = title.add_run("ABSTRACT\n")
            runAbstract.bold = True

            abstractText = document.add_paragraph(
                "Write here...")
            abstractText.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            abstractText.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
            abstractText.paragraph_format.space_before = 0
            abstractText.paragraph_format.space_after = 0
            abstractText.paragraph_format.first_line_indent = Cm(1.25)

            pageCount += 1
            document.add_page_break()


def Summary():
    global pageCount
    if config['summary'] == 1:
        title = document.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        title.paragraph_format.space_before = 0
        title.paragraph_format.space_after = 0
        runResume = title.add_run("SUMÁRIO\n")
        runResume.bold = True

        summary = document.add_paragraph()
        summary.alignment = WD_ALIGN_PARAGRAPH.LEFT
        summary.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        summary.paragraph_format.space_before = 0
        summary.paragraph_format.space_after = 0
        tab_stops = summary.paragraph_format.tab_stops
        tab_stop = tab_stops.add_tab_stop(
            Cm(15.5), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS)

        pageCount += 1
        if pageCount < 10:
            zero = '0'
        else:
            zero = ''
        runIntro = summary.add_run("1 INTRODUÇÃO\t"+zero+str(pageCount)+'\n')
        runIntro.bold = True

        pageCount += 1
        if pageCount < 10:
            zero = '0'
        else:
            zero = ''
        runDev = summary.add_run(
            "2 DESENVOLVIMENTO\t"+zero+str(pageCount)+'\n')
        runDev.bold = True

        pageCount += 1
        if pageCount < 10:
            zero = '0'
        else:
            zero = ''
        runConc = summary.add_run("3 CONCLUSÃO\t"+zero+str(pageCount)+'\n')
        runConc.bold = True

        pageCount += 1
        if pageCount < 10:
            zero = '0'
        else:
            zero = ''
        runRef = summary.add_run("REFERÊNCIAS\t"+zero+str(pageCount)+'\n')
        runRef.bold = True

        document.add_page_break()


def Introduction():
    title = document.add_paragraph()
    title.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    title.paragraph_format.space_before = 0
    title.paragraph_format.space_after = 0
    runResume = title.add_run("1 INTRODUÇÃO\n")
    runResume.bold = True

    intro = document.add_paragraph("Write here...")
    intro = document.add_paragraph()
    intro = document.add_paragraph()
    intro = document.add_paragraph()
    intro = document.add_paragraph()
    intro.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    intro.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    intro.paragraph_format.space_before = 0
    intro.paragraph_format.space_after = 0
    intro.paragraph_format.first_line_indent = Cm(1.25)

    document.add_page_break()


def Deployment():
    title = document.add_paragraph()
    title.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    title.paragraph_format.space_before = 0
    title.paragraph_format.space_after = 0
    runResume = title.add_run("2 DESENVOLVIMENTO\n")
    runResume.bold = True

    dev = document.add_paragraph(
        "Write here...")
    dev = document.add_paragraph()
    dev = document.add_paragraph()
    dev = document.add_paragraph()
    dev = document.add_paragraph()
    dev = document.add_paragraph()
    dev = document.add_paragraph()
    dev.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    dev.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    dev.paragraph_format.space_before = 0
    dev.paragraph_format.space_after = 0
    dev.paragraph_format.first_line_indent = Cm(1.25)

    document.add_page_break()


def Conclusion():
    title = document.add_paragraph()
    title.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    title.paragraph_format.space_before = 0
    title.paragraph_format.space_after = 0
    runResume = title.add_run("3 CONCLUSÃO\n")
    runResume.bold = True

    conclusion = document.add_paragraph("Write Here...")
    conclusion = document.add_paragraph()
    conclusion = document.add_paragraph()
    conclusion = document.add_paragraph()
    conclusion = document.add_paragraph()
    conclusion.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    conclusion.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    conclusion.paragraph_format.space_before = 0
    conclusion.paragraph_format.space_after = 0
    conclusion.paragraph_format.first_line_indent = Cm(1.25)

    document.add_page_break()


def References():
    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    title.paragraph_format.space_before = 0
    title.paragraph_format.space_after = 0
    runAbstract = title.add_run("REFERÊNCIAS\n\n\n")
    runAbstract.bold = True


Cover()
BackCover()
Dedicatory()
Thanks()
Epigraph()
Resume()
Summary()
Introduction()
Deployment()
Conclusion()
References()

document.save('demo.docx')

'''
.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
.paragraph_format.space_before = 0
.paragraph_format.space_after = 0
.paragraph_format.first_line_indent = Cm(1.25)
'''
